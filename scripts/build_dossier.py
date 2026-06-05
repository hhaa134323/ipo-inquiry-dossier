#!/usr/bin/env python3
"""Build a .docx dossier from hits.jsonl + .txt cache + source PDFs.

Usage:
    python scripts/build_dossier.py
    python scripts/build_dossier.py --input ../my_pdfs --output ../out
    python scripts/build_dossier.py --hits ../custom_hits.jsonl

Reads hits.jsonl (first line = meta, rest = hits), extracts content according
to each hit's page-range and form-type spec, and writes a .docx dossier.

最终 .docx 写入 --output 目录；所有中间产物（截图、表格截图、默认 hits.jsonl）
统一落到 --work 目录（默认 {output}/_work），使 --output 只保留最终底稿。
"""

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz  # pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

# ---------------------------------------------------------------------------
#  CJK font helpers (fix □ rendering of simplified Chinese in Word)
# ---------------------------------------------------------------------------


def set_cjk_font(doc, name="宋体"):
    """Set CJK east-asia font on the Normal style so Chinese chars render."""
    style = doc.styles["Normal"]
    style.font.name = name
    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(slot), name)


def set_run_cjk(run, name="宋体"):
    """Ensure a specific run uses the CJK eastAsia font (fallback per run)."""
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def _apply_cjk_to_all(doc, name="宋体"):
    """Iterate every run in the document body and table cells — safety net."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_cjk(run, name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_cjk(run, name)


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

_SENTENCE_ENDS = re.compile(r'[。？！；）)》。"]$')


def add_toc(doc: Document) -> None:
    """Insert a Word TOC field (headings 1-3, hyperlinked, no tab leader)."""
    paragraph = doc.add_paragraph()

    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z '
    run2._element.append(instr)

    run._element.append(OxmlElement("w:r"))
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_sep)

    run3 = paragraph.add_run("[目录将在 Word 中右键 -> 更新域后显示]")

    run._element.append(OxmlElement("w:r"))
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)


def clean_pdf_text(text: str) -> str:
    """Clean PDF-extracted text for clean copy-paste.

    - Removes standalone page-number lines.
    - Merges lines broken by PDF layout (soft line-breaks).
    - Preserves paragraph breaks (double newlines).
    """
    lines = text.splitlines()
    cleaned: list[str] = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\d{1,4}", stripped):
            continue
        if re.fullmatch(r"\d{1,3}[-u2013]\d{1,3}", stripped):
            continue
        cleaned.append(stripped)

    merged: list[str] = []
    for line in cleaned:
        if not line:
            merged.append("")
        elif merged and merged[-1] and not _SENTENCE_ENDS.search(merged[-1]):
            merged[-1] += line
        else:
            merged.append(line)

    result: list[str] = []
    blank = False
    for line in merged:
        if not line:
            if not blank:
                result.append("")
                blank = True
        else:
            result.append(line)
            blank = False

    return "\n".join(result)


def add_para_with_anchors(doc, text: str, anchors: list[str]) -> None:
    """Add a paragraph; any anchor substring found verbatim gets yellow highlight.

    Anchors are expected to be short (8-30 chars) verbatim excerpts from the
    original PDF text.  If an anchor cannot be found (e.g. whitespace mismatch
    after PDF cleaning) it is silently skipped — never approximate text.
    """
    if not anchors:
        doc.add_paragraph(text)
        return
    unique = sorted(set(a.strip() for a in anchors if a.strip()),
                    key=len, reverse=True)
    if not unique:
        doc.add_paragraph(text)
        return

    # Mark character-level highlight regions
    highlight = [False] * len(text)
    for anchor in unique:
        start = 0
        while True:
            pos = text.find(anchor, start)
            if pos == -1:
                break
            for i in range(pos, pos + len(anchor)):
                if i < len(text):
                    highlight[i] = True
            start = pos + 1

    # Build runs: contiguous highlighted / non-highlighted segments
    p = doc.add_paragraph()
    i = 0
    while i < len(text):
        j = i
        in_hl = highlight[i]
        while j < len(text) and highlight[j] == in_hl:
            j += 1
        run = p.add_run(text[i:j])
        if in_hl:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        i = j


def extract_text_by_range(txt_path: Path, start_page: int,
                          end_page: int) -> str:
    """Extract and clean text from .txt cache for a page range (1-indexed)."""
    if not txt_path.exists():
        return f"[缓存文件 {txt_path.name} 不存在]"

    text = txt_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    in_range = False
    captured: list[str] = []
    for line in lines:
        if line.startswith(f"[PAGE {start_page}]"):
            in_range = True
        elif line.startswith(f"[PAGE {end_page + 1}]"):
            break
        if in_range:
            captured.append(line)

    captured = [l for l in captured if not l.startswith("[PAGE ")]
    return clean_pdf_text("\n".join(captured))


def extract_table_from_pdf(pdf_path: Path,
                           page_num: int) -> list[list[str]] | None:
    """Extract the first table from a PDF page using pymupdf find_tables()."""
    try:
        doc = fitz.open(pdf_path)
    except Exception:
        return None
    if page_num < 1 or page_num > len(doc):
        doc.close()
        return None
    page = doc[page_num - 1]
    tables = page.find_tables()
    doc.close()
    if tables.tables:
        raw = tables.tables[0].extract()
        return [[str(c) if c is not None else "" for c in row]
                for row in raw]
    return None


def render_page_image(pdf_path: Path, page_num: int,
                      work_dir: Path) -> Path | None:
    """Render a PDF page as a 2x PNG, return the saved path."""
    try:
        doc = fitz.open(pdf_path)
    except Exception:
        return None
    if page_num < 1 or page_num > len(doc):
        doc.close()
        return None
    page = doc[page_num - 1]
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img_name = f"{pdf_path.stem}_p{page_num}.png"
    img_path = work_dir / img_name
    pix.save(str(img_path))
    doc.close()
    return img_path


def add_table_to_doc(doc: Document, data: list[list[str]]) -> None:
    """Add a 2D table to the document with Table Grid style."""
    if not data:
        return
    rows = len(data)
    cols = max(len(r) for r in data) if data else 0
    if cols == 0:
        return
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            t.rows[i].cells[j].text = str(cell_text)


# ---------------------------------------------------------------------------
#  Table-aware page rendering (表格几何抽表 + 三级兜底)
# ---------------------------------------------------------------------------


def _get_table_health(data: list[list[str]]) -> tuple[bool, int, int, float]:
    """Evaluate table health.

    Healthy if rows≥2, cols≥2, and empty-cell ratio < 50%.
    Returns (is_healthy, rows, cols, empty_ratio).
    """
    if not data:
        return False, 0, 0, 1.0
    rows = len(data)
    cols = max(len(r) for r in data) if data else 0
    if rows < 2 or cols < 2:
        return False, rows, cols, 1.0
    total = rows * cols
    empty = sum(
        1 for r in data
        for c in r
        if not str(c).strip()
    )
    ratio = empty / total if total > 0 else 1.0
    return ratio < 0.5, rows, cols, ratio


def _clip_page_region(pdf_path: Path, page_num: int,
                      bbox: tuple[float, float, float, float],
                      work_dir: Path) -> Path | None:
    """Clip a page region as 2× PNG; return saved path or None.

    *bbox* is (x0, y0, x1, y1) in PDF points, as returned by PyMuPDF
    ``find_tables()`` (plain tuple, not fitz.Rect).  PNG is saved to a
    dedicated subdirectory under work_dir.
    """
    try:
        pdf = fitz.open(pdf_path)
    except Exception:
        return None
    if page_num < 1 or page_num > len(pdf):
        pdf.close()
        return None
    page = pdf[page_num - 1]
    pix = page.get_pixmap(clip=bbox, matrix=fitz.Matrix(2, 2))
    clip_dir = work_dir / "_table_clips"
    clip_dir.mkdir(parents=True, exist_ok=True)
    img_name = f"{pdf_path.stem}_p{page_num}_table.png"
    img_path = clip_dir / img_name
    pix.save(str(img_path))
    pdf.close()
    return img_path


def _render_page_interleaved(doc: Document, pdf_path: Path,
                             work_dir: Path, page_num: int,
                             anchors: list[str] | None) -> None:
    """Render one PDF page with interleaved text blocks and tables.

    Three-level fallback per detected table:
      1. Healthy (rows≥2, cols≥2, empty<50%) → python-docx Table w/ Grid style
      2. Unhealthy → 2× PNG clip of the table bbox embedded as picture
      3. No table on page → normal paragraph text only

    Text blocks whose bounding boxes fall inside a detected table are
    skipped to avoid duplicating table content as plain text.  Yellow
    highlight (anchors) continues to apply to text paragraphs only.
    """
    try:
        pdf = fitz.open(pdf_path)
    except Exception:
        return
    if page_num < 1 or page_num > len(pdf):
        pdf.close()
        return

    page = pdf[page_num - 1]

    # ── 1. Detect tables ──────────────────────────────────────────────
    table_finder = page.find_tables()
    table_items: list[dict[str, Any]] = []
    for t in table_finder.tables:
        bbox = t.bbox
        raw = t.extract()
        if not raw:
            continue
        str_data = [[str(c) if c is not None else "" for c in row]
                     for row in raw]
        is_healthy, *_ = _get_table_health(str_data)
        table_items.append({
            "y0": bbox[1],
            "type": "table",
            "bbox": bbox,
            "data": str_data,
            "is_healthy": is_healthy,
        })

    # ── 2. Get text blocks, skip those inside table bboxes ───────────
    text_items: list[dict[str, Any]] = []
    for block in page.get_text("blocks"):
        if len(block) < 6:
            continue
        if not isinstance(block[4], str):
            continue          # skip image / non-text blocks
        raw_text = block[4].strip()
        if not raw_text:
            continue
        x0, y0, x1, y1 = block[0], block[1], block[2], block[3]

        # Skip if the centre of this block lies inside a detected table
        cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
        inside_table = False
        for tbl in table_items:
            tb = tbl["bbox"]
            if (tb[0] - 2 <= cx <= tb[2] + 2 and
                    tb[1] - 2 <= cy <= tb[3] + 2):
                inside_table = True
                break
        if inside_table:
            continue

        text_items.append({
            "y0": y0,
            "type": "text",
            "text": raw_text,
        })

    pdf.close()

    # ── 3. Merge & sort by vertical reading order ─────────────────────
    items = sorted(text_items + table_items, key=lambda x: x["y0"])

    # ── 4. Render in order ────────────────────────────────────────────
    for item in items:
        if item["type"] == "text":
            cleaned = clean_pdf_text(item["text"])
            add_para_with_anchors(doc, cleaned, anchors or [])
        else:
            if item["is_healthy"]:
                add_table_to_doc(doc, item["data"])
            else:
                img = _clip_page_region(pdf_path, page_num,
                                        item["bbox"], work_dir)
                if img:
                    doc.add_paragraph("（表格区域截图）")
                    doc.add_picture(str(img), width=Inches(6))


def render_form(doc: Document, pdf_path: Path, txt_path: Path,
                work_dir: Path,
                page_num: int, form: str, *, is_range: bool = False,
                end_page: int | None = None,
                anchors: list[str] | None = None) -> None:
    """Render content for a given page/range.

    * form='图'   — unchanged: full-page PNG per page.
    * form='表格' — auto table-detect + interleaved rendering per page.
    * form='文本' — auto table-detect + interleaved rendering per page
                     (replaces old .txt-cache plain-text dump).

    Tables are detected via PyMuPDF ``find_tables()`` on each source page
    and rendered as python-docx Tables (or PNG clip fallback for unhealthy
    tables).  Text blocks outside table bboxes yield normal paragraphs
    with yellow-highlight anchor support.
    """
    if form == "图":
        pages = range(page_num, (end_page or page_num) + 1)
        for p in pages:
            img_path = render_page_image(pdf_path, p, work_dir)
            if img_path:
                doc.add_paragraph(f"（第{p}页截图）")
                doc.add_picture(str(img_path), width=Inches(6))
    else:
        # 文本/表格：自动检测表格，交错渲染，不依赖 .txt 缓存
        pages = range(page_num, (end_page or page_num) + 1)
        for p in pages:
            _render_page_interleaved(doc, pdf_path, work_dir, p, anchors)


# ---------------------------------------------------------------------------
#  Main
# ---------------------------------------------------------------------------

def build_dossier() -> None:
    """Main entry point: read hits.jsonl (meta + hits), generate .docx."""
    parser = argparse.ArgumentParser(
        description="从 hits.jsonl 生成 IPO 问询可比案例底稿 .docx"
    )
    parser.add_argument(
        "--input", "-i",
        default="./input",
        help="输入目录（含 PDF 与 .txt 缓存），默认 ./input",
    )
    parser.add_argument(
        "--output", "-o",
        default="./output",
        help="输出目录（只放最终 .docx），默认 ./output",
    )
    parser.add_argument(
        "--work",
        default=None,
        help="中间产物目录（截图、表格截图、默认 hits.jsonl 等），默认 {output}/_work",
    )
    parser.add_argument(
        "--hits",
        default=None,
        help="hits.jsonl 路径（默认 {work}/hits.jsonl）",
    )
    args = parser.parse_args()

    input_dir = Path(args.input).resolve()
    output_dir = Path(args.output).resolve()
    work_dir = Path(args.work).resolve() if args.work else output_dir / "_work"
    hits_path = Path(args.hits).resolve() if args.hits else work_dir / "hits.jsonl"

    output_dir.mkdir(parents=True, exist_ok=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    if not hits_path.exists():
        print(f"ERROR: {hits_path} not found."
              " Run the retrieval step first (see METHODOLOGY.md step 5).")
        return

    # --- Read hits.jsonl: first line = meta, rest = hit entries ---
    lines = hits_path.read_text(encoding="utf-8").strip().splitlines()
    if not lines:
        print("ERROR: hits.jsonl is empty.")
        return

    # Parse meta (first line)
    try:
        meta: dict[str, Any] = json.loads(lines[0])
    except json.JSONDecodeError as e:
        print(f"ERROR: first line of hits.jsonl must be a valid JSON meta object — {e}")
        return

    # Parse hits (remaining lines)
    hits: list[dict[str, Any]] = []
    for line in lines[1:]:
        stripped = line.strip()
        if stripped:
            try:
                hits.append(json.loads(stripped))
            except json.JSONDecodeError as e:
                print(f"WARNING: skipping invalid JSON line: {e}")

    if not hits:
        print("ERROR: hits.jsonl has meta but no hit entries.")
        return

    # --- Filter by 判可比依据.保留 (精排保留标志) ---
    kept_hits = [h for h in hits if h.get("判可比依据", {}).get("保留") is True]
    excluded_hits = [h for h in hits if h not in kept_hits]
    if not kept_hits:
        print("WARNING: no hit entries marked '保留:true'. Generating empty dossier.")
        # Still continue — we'll generate the file with meta + excluded appendix only

    topic = meta.get("主题", "可比案例")
    question_text = meta.get("问题原文", "")
    search_terms = meta.get("检索条件", "")

    # Derive date from first kept hit or today
    date_str = (kept_hits[0].get("日期", datetime.now().strftime("%Y%m%d"))
                if kept_hits else datetime.now().strftime("%Y%m%d"))

    safe_topic = re.sub(r'[\\/:*?"<>|]', "_", str(topic))
    safe_date = re.sub(r'[\\/:*?"<>|]', "_", str(date_str))

    doc = Document()
    set_cjk_font(doc)  # Fix CJK rendering

    # ---- Title ----
    title = doc.add_heading("IPO 问询可比案例底稿", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ---- 问题原文 (verbatim, full) ----
    doc.add_heading("问询问题原文", level=2)
    doc.add_paragraph(question_text)

    doc.add_paragraph("")

    # ---- 基本信息 ----
    doc.add_heading("基本信息", level=2)
    doc.add_paragraph(f"主题：{topic}")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"保留案例：{len(kept_hits)} / 总候选：{len(hits)}")
    if excluded_hits:
        doc.add_paragraph(f"已排除候选：{len(excluded_hits)} 个（见文末附录）")
    if search_terms:
        doc.add_paragraph(f"检索条件：{search_terms}")
    doc.add_paragraph("")

    # ---- 结论速览卡 ----
    doc.add_heading("结论速览卡", level=1)
    if kept_hits:
        # Build summary table
        summary_rows = [["问询主题", topic]]
        for h in kept_hits:
            co = h.get("公司", "")
            score = h.get("判可比依据", {}).get("总分", "?")
            gaps = h.get("判可比依据", {}).get("缺口说明", "")
            relevance = h.get("相关性", "")
            summary_rows.append(["保留案例", f"{co}（精排总分 {score}/10）"])
            summary_rows.append(["精排结论", relevance[:120] + ("…" if len(relevance) > 120 else "")])
            summary_rows.append(["缺口说明", gaps[:120] + ("…" if len(gaps) > 120 else "")])
        # Extract transplantable dimensions from scoring
        transplant = []
        if kept_hits and kept_hits[0].get("判可比依据", {}).get("评分"):
            for s in kept_hits[0]["判可比依据"]["评分"]:
                if s.get("得分", 0) >= 2:
                    transplant.append(s.get("维度", ""))
        if transplant:
            summary_rows.append(["可直接移植章节", "、".join(transplant)])
        tbl = doc.add_table(rows=len(summary_rows), cols=2)
        tbl.style = "Light Shading Accent 1"
        for i, (k, v) in enumerate(summary_rows):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = str(v)
        doc.add_paragraph("")
    else:
        doc.add_paragraph("⚠ 本次检索未找到精排保留的合格可比案例。")
        doc.add_paragraph("")

    # ---- TOC ----
    doc.add_heading("目录", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ---- Each case (only kept_hits) ----
    for hit in kept_hits:
        seq = hit.get("序号", "")
        company = hit.get("公司", "未知公司")
        round_ = hit.get("轮次", "")
        date = hit.get("日期", "")
        filename = hit.get("文件", "")
        inquiry = hit.get("问询", {})
        reply = hit.get("回复", {})
        relevance = hit.get("相关性", "")

        inquiry_page = inquiry.get("页", 0)
        inquiry_form = inquiry.get("形式", "文本")
        reply_start = reply.get("起页", 0)
        reply_end = reply.get("止页", reply_start)
        reply_form = reply.get("形式", "文本")

        try:
            inquiry_page = int(inquiry_page)
        except (ValueError, TypeError):
            inquiry_page = 0
        try:
            reply_start = int(reply_start)
        except (ValueError, TypeError):
            reply_start = 0
        try:
            reply_end = int(reply_end) if reply_end else reply_start
        except (ValueError, TypeError):
            reply_end = reply_start

        pdf_path = input_dir / filename
        txt_path = pdf_path.with_suffix(".txt")

        # ---- Case heading ----
        heading_text = f"案例{seq}：{company}"
        if date:
            heading_text += f"（{date}）"
        doc.add_heading(heading_text, level=1)

        # ---- Meta table with 5-level provenance ----
        provenance_rows = [
            ("公司", company),
            ("轮次", round_),
            ("日期", date),
            ("文件", filename),
            ("问询页码", f"第 {inquiry_page} 页" if inquiry_page > 0 else "无"),
            ("回复页码", f"第 {reply_start}-{reply_end} 页" if reply_end != reply_start and reply_start > 0
                        else f"第 {reply_start} 页" if reply_start > 0 else "无"),
            ("精排总分", f'{hit.get("判可比依据", {}).get("总分", "?")}/10'),
            ("相关性", relevance),
        ]
        tbl = doc.add_table(rows=len(provenance_rows), cols=2)
        tbl.style = "Light Shading Accent 1"
        for i, (k, v) in enumerate(provenance_rows):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = str(v)
        doc.add_paragraph("")

        # ---- Extract anchor phrases for yellow-highlight ----
        anchors_map = hit.get("关键锚点", {}) or {}
        inquiry_anchors = anchors_map.get("问询", [])
        reply_anchors = anchors_map.get("回复", [])

        # ---- Inquiry section ----
        if inquiry_page > 0:
            doc.add_heading(f"问询（第 {inquiry_page} 页）", level=2)
            render_form(doc, pdf_path, txt_path, work_dir,
                        inquiry_page, inquiry_form,
                        anchors=inquiry_anchors)

        # ---- Reply section ----
        if reply_start > 0:
            range_str = (
                f"第 {reply_start}-{reply_end} 页"
                if reply_end != reply_start
                else f"第 {reply_start} 页"
            )
            doc.add_heading(f"回复（{range_str}）", level=2)
            render_form(
                doc, pdf_path, txt_path, work_dir,
                reply_start, reply_form,
                is_range=(reply_end != reply_start),
                end_page=reply_end,
                anchors=reply_anchors,
            )

        doc.add_page_break()

    # ---- Excluded cases appendix ----
    if excluded_hits:
        doc.add_heading("附录：已排除候选", level=1)
        doc.add_paragraph("以下候选经精排判定未达保留标准，未进入正文。")
        for h in excluded_hits:
            co = h.get("公司", "未知公司")
            seq = h.get("序号", "")
            score = h.get("判可比依据", {}).get("总分", "?")
            reason = h.get("判可比依据", {}).get("缺口说明", "") or h.get("丢弃原因", "未说明")
            doc.add_heading(f"排除案例{seq}：{co}（总分 {score}/10）", level=3)
            p = doc.add_paragraph()
            run = p.add_run(f"排除原因：{reason}")
            run.bold = True
            doc.add_paragraph("")

    # ---- Apply CJK font to every run as safety net ----
    _apply_cjk_to_all(doc)

    # ---- Save ----
    out_name = f"底稿_{safe_topic}_{safe_date}.docx"
    out_path = output_dir / out_name
    doc.save(str(out_path))
    print(f"Dossier saved to: {out_path}")


if __name__ == "__main__":
    build_dossier()
